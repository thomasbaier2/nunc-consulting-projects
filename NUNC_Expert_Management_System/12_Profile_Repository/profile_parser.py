#!/usr/bin/env python3
"""
NUNC Expert Management System - Profile Parser
Parst Word/PDF Profile und konvertiert sie zu Supabase-Format
"""

import os
import re
import json
from datetime import datetime, date
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, asdict
from pathlib import Path

# PDF Processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word Processing
try:
    from docx import Document
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

@dataclass
class ParsedProfile:
    """Geparstes Profil aus Word/PDF"""
    file_path: str
    file_type: str  # 'word' oder 'pdf'
    
    # Grunddaten
    first_name: str = ""
    last_name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    
    # Skills und Erfahrung
    technical_skills: List[str] = None
    soft_skills: List[str] = None
    certifications: List[str] = None
    languages: List[str] = None
    
    # Berufserfahrung
    work_experience: List[Dict] = None
    education: List[Dict] = None
    
    # Projekt-Erfahrung (für Cross-connections)
    company_experience: Dict[str, str] = None
    project_experience: Dict[str, str] = None
    industry_experience: List[str] = None
    
    # Verfügbarkeit
    availability_status: str = "unknown"
    next_available_date: Optional[date] = None
    preferred_hours_per_week: Optional[int] = None
    remote_preference: str = "flexible"
    
    # Metadaten
    raw_text: str = ""
    parsing_confidence: float = 0.0
    parsing_errors: List[str] = None
    
    def __post_init__(self):
        if self.technical_skills is None:
            self.technical_skills = []
        if self.soft_skills is None:
            self.soft_skills = []
        if self.certifications is None:
            self.certifications = []
        if self.languages is None:
            self.languages = []
        if self.work_experience is None:
            self.work_experience = []
        if self.education is None:
            self.education = []
        if self.company_experience is None:
            self.company_experience = {}
        if self.project_experience is None:
            self.project_experience = {}
        if self.industry_experience is None:
            self.industry_experience = []
        if self.parsing_errors is None:
            self.parsing_errors = []

class ProfileParser:
    """Parst Word/PDF Profile und extrahiert strukturierte Daten"""
    
    def __init__(self):
        """Initialisiert den ProfileParser"""
        self.technical_skills_keywords = [
            'python', 'java', 'javascript', 'react', 'angular', 'vue', 'node.js',
            'aws', 'azure', 'docker', 'kubernetes', 'sql', 'postgresql', 'mongodb',
            'spring', 'django', 'flask', 'express', 'typescript', 'html', 'css',
            'git', 'jenkins', 'terraform', 'ansible', 'linux', 'windows', 'macos'
        ]
        
        self.soft_skills_keywords = [
            'teamwork', 'leadership', 'communication', 'problem-solving',
            'analytical', 'creative', 'detail-oriented', 'time-management',
            'project-management', 'agile', 'scrum', 'mentoring'
        ]
        
        self.company_keywords = [
            'bmw', 'mercedes', 'audi', 'volkswagen', 'sap', 'microsoft', 'google',
            'amazon', 'apple', 'facebook', 'meta', 'tesla', 'siemens', 'bosch'
        ]
    
    def parse_file(self, file_path: str) -> ParsedProfile:
        """Parst eine Datei (Word oder PDF)"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Bestimme Dateityp
        if file_path.suffix.lower() == '.docx':
            return self._parse_word_file(str(file_path))
        elif file_path.suffix.lower() == '.pdf':
            return self._parse_pdf_file(str(file_path))
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
    
    def _parse_word_file(self, file_path: str) -> ParsedProfile:
        """Parst eine Word-Datei"""
        if not WORD_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        try:
            doc = Document(file_path)
            
            # Extrahiere Text aus allen Absätzen
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extrahiere Text aus Tabellen
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            raw_text = '\n'.join(text_parts)
            
            # Erstelle ParsedProfile
            profile = ParsedProfile(
                file_path=file_path,
                file_type='word',
                raw_text=raw_text
            )
            
            # Parse den Text
            self._parse_text_content(profile)
            
            return profile
            
        except Exception as e:
            raise Exception(f"Error parsing Word file {file_path}: {e}")
    
    def _parse_pdf_file(self, file_path: str) -> ParsedProfile:
        """Parst eine PDF-Datei"""
        if not PDF_AVAILABLE:
            raise ImportError("PDF libraries not installed. Run: pip install PyPDF2 pdfplumber")
        
        try:
            # Versuche zuerst pdfplumber (besser für Text-Extraktion)
            try:
                with pdfplumber.open(file_path) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    raw_text = '\n'.join(text_parts)
            except:
                # Fallback zu PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_parts = []
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    raw_text = '\n'.join(text_parts)
            
            # Erstelle ParsedProfile
            profile = ParsedProfile(
                file_path=file_path,
                file_type='pdf',
                raw_text=raw_text
            )
            
            # Parse den Text
            self._parse_text_content(profile)
            
            return profile
            
        except Exception as e:
            raise Exception(f"Error parsing PDF file {file_path}: {e}")
    
    def _parse_text_content(self, profile: ParsedProfile):
        """Parst den extrahierten Text und extrahiert strukturierte Daten"""
        text = profile.raw_text.lower()
        
        # Extrahiere Namen aus Dateinamen falls nicht im Text gefunden
        if not profile.first_name or not profile.last_name:
            self._extract_name_from_filename(profile)
        
        # Extrahiere E-Mail (optional - kann leer bleiben)
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', profile.raw_text)
        if email_match:
            profile.email = email_match.group()
        else:
            # Generiere E-Mail aus Namen falls keine gefunden
            if profile.first_name and profile.last_name:
                profile.email = f"{profile.first_name.lower()}.{profile.last_name.lower()}@nunc.local"
        
        # Extrahiere Telefon
        phone_patterns = [
            r'\+49\s*\d{2,4}\s*\d{3,4}\s*\d{3,4}',
            r'0\d{2,4}\s*\d{3,4}\s*\d{3,4}',
            r'\(\d{2,4}\)\s*\d{3,4}\s*\d{3,4}'
        ]
        for pattern in phone_patterns:
            phone_match = re.search(pattern, profile.raw_text)
            if phone_match:
                profile.phone = phone_match.group()
                break
        
        # Extrahiere Namen (erste Zeile oder nach "Name:")
        name_patterns = [
            r'Name:\s*([A-Za-z\s]+)',
            r'^([A-Za-z\s]+)$',
            r'([A-Za-z]+)\s+([A-Za-z]+)'
        ]
        
        for pattern in name_patterns:
            name_match = re.search(pattern, profile.raw_text, re.MULTILINE)
            if name_match:
                if len(name_match.groups()) == 2:
                    profile.first_name = name_match.group(1).strip()
                    profile.last_name = name_match.group(2).strip()
                else:
                    full_name = name_match.group(1).strip()
                    name_parts = full_name.split()
                    if len(name_parts) >= 2:
                        profile.first_name = name_parts[0]
                        profile.last_name = ' '.join(name_parts[1:])
                break
        
        # Extrahiere technische Skills
        profile.technical_skills = self._extract_skills(text, self.technical_skills_keywords)
        
        # Extrahiere Soft Skills
        profile.soft_skills = self._extract_skills(text, self.soft_skills_keywords)
        
        # Extrahiere Zertifizierungen
        profile.certifications = self._extract_certifications(text)
        
        # Extrahiere Sprachen
        profile.languages = self._extract_languages(text)
        
        # Extrahiere Berufserfahrung
        profile.work_experience = self._extract_work_experience(text)
        
        # Extrahiere Firmen-Erfahrung für Cross-connections
        profile.company_experience = self._extract_company_experience(text)
        
        # Extrahiere Branchen-Erfahrung
        profile.industry_experience = self._extract_industry_experience(text)
        
        # Berechne Parsing-Confidence
        profile.parsing_confidence = self._calculate_parsing_confidence(profile)
    
    def _extract_skills(self, text: str, keywords: List[str]) -> List[str]:
        """Extrahiert Skills aus dem Text"""
        found_skills = []
        text_lower = text.lower()
        
        for skill in keywords:
            if skill.lower() in text_lower:
                found_skills.append(skill.title())
        
        return list(set(found_skills))  # Entferne Duplikate
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extrahiert Zertifizierungen"""
        cert_patterns = [
            r'(AWS|Azure|Google Cloud|Microsoft|Oracle|Cisco|PMP|ITIL|Agile|Scrum)\s+[A-Za-z\s]+',
            r'Certified\s+[A-Za-z\s]+',
            r'[A-Za-z]+\s+Certification'
        ]
        
        certifications = []
        for pattern in cert_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            certifications.extend(matches)
        
        return list(set(certifications))
    
    def _extract_languages(self, text: str) -> List[str]:
        """Extrahiert Sprachen"""
        language_patterns = [
            r'(German|English|French|Spanish|Italian|Portuguese|Dutch|Russian|Chinese|Japanese|Korean)',
            r'(Deutsch|Englisch|Französisch|Spanisch|Italienisch|Portugiesisch|Niederländisch|Russisch|Chinesisch|Japanisch|Koreanisch)'
        ]
        
        languages = []
        for pattern in language_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            languages.extend(matches)
        
        return list(set(languages))
    
    def _extract_work_experience(self, text: str) -> List[Dict]:
        """Extrahiert Berufserfahrung"""
        # Vereinfachte Extraktion - kann erweitert werden
        experience = []
        
        # Suche nach Firmen-Namen
        company_pattern = r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)'
        companies = re.findall(company_pattern, text)
        
        for company in companies[:5]:  # Max. 5 Firmen
            if len(company) > 3:  # Mindestlänge für Firmenname
                experience.append({
                    'company': company,
                    'position': 'Unknown',
                    'duration': 'Unknown',
                    'description': ''
                })
        
        return experience
    
    def _extract_company_experience(self, text: str) -> Dict[str, str]:
        """Extrahiert Firmen-Erfahrung für Cross-connections"""
        company_experience = {}
        text_lower = text.lower()
        
        for company in self.company_keywords:
            if company in text_lower:
                # Versuche Dauer zu extrahieren
                duration_pattern = rf'{company}[^.]*?(\d+)\s*(year|jahr|jahren)'
                duration_match = re.search(duration_pattern, text_lower)
                
                if duration_match:
                    company_experience[company.title()] = f"{duration_match.group(1)} years"
                else:
                    company_experience[company.title()] = "Experience found"
        
        return company_experience
    
    def _extract_industry_experience(self, text: str) -> List[str]:
        """Extrahiert Branchen-Erfahrung"""
        industries = [
            'automotive', 'finance', 'healthcare', 'retail', 'manufacturing',
            'consulting', 'technology', 'telecommunications', 'energy', 'logistics'
        ]
        
        found_industries = []
        text_lower = text.lower()
        
        for industry in industries:
            if industry in text_lower:
                found_industries.append(industry.title())
        
        return found_industries
    
    def _calculate_parsing_confidence(self, profile: ParsedProfile) -> float:
        """Berechnet die Parsing-Confidence (0.0 - 1.0)"""
        confidence = 0.0
        
        # Grunddaten (40% Gewichtung)
        if profile.first_name and profile.last_name:
            confidence += 0.2
        if profile.email:
            confidence += 0.1
        if profile.phone:
            confidence += 0.1
        
        # Skills (30% Gewichtung)
        if profile.technical_skills:
            confidence += 0.15
        if profile.soft_skills:
            confidence += 0.15
        
        # Erfahrung (30% Gewichtung)
        if profile.work_experience:
            confidence += 0.15
        if profile.company_experience:
            confidence += 0.15
        
        return min(confidence, 1.0)
    
    def _extract_name_from_filename(self, profile: ParsedProfile):
        """Extrahiert Namen aus dem Dateinamen"""
        filename = Path(profile.file_path).stem  # Ohne .docx/.pdf
        
        # Entferne NUNC Profile Präfix
        filename = re.sub(r'^\d+_NUNC\s+Profile\s+', '', filename, flags=re.IGNORECASE)
        filename = re.sub(r'^\d+_NUNC\s+Profil\s+', '', filename, flags=re.IGNORECASE)
        filename = re.sub(r'^NUNC\s+Profile\s+', '', filename, flags=re.IGNORECASE)
        filename = re.sub(r'^NUNC\s+Profil\s+', '', filename, flags=re.IGNORECASE)
        
        # Entferne Rollen/Positionen am Ende
        filename = re.sub(r'\s+(Salesforce|Mulesoft|Tableau|Project|Data|Commerce|Berater|Consultant|Architect|Developer|Expert|Manager|Marketer|Scientist).*$', '', filename, flags=re.IGNORECASE)
        
        # Teile in Vor- und Nachname
        name_parts = filename.strip().split()
        
        if len(name_parts) >= 2:
            profile.first_name = name_parts[0]
            profile.last_name = ' '.join(name_parts[1:])
        elif len(name_parts) == 1:
            profile.first_name = name_parts[0]
            profile.last_name = "Unknown"
        else:
            profile.first_name = "Unknown"
            profile.last_name = "Unknown"
    
    def parse_directory(self, directory_path: str) -> List[ParsedProfile]:
        """Parst alle Profile in einem Verzeichnis"""
        directory = Path(directory_path)
        profiles = []
        
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        # Suche nach Word und PDF Dateien
        for file_path in directory.rglob('*.docx'):
            try:
                profile = self.parse_file(str(file_path))
                profiles.append(profile)
                print(f"Parsed: {file_path.name}")
            except Exception as e:
                print(f"Error parsing {file_path.name}: {e}")
        
        for file_path in directory.rglob('*.pdf'):
            try:
                profile = self.parse_file(str(file_path))
                profiles.append(profile)
                print(f"Parsed: {file_path.name}")
            except Exception as e:
                print(f"Error parsing {file_path.name}: {e}")
        
        return profiles

# Beispiel-Verwendung
if __name__ == "__main__":
    parser = ProfileParser()
    
    # Test mit einem einzelnen File
    # profile = parser.parse_file("path/to/profile.docx")
    # print(f"Parsed profile: {profile.first_name} {profile.last_name}")
    # print(f"Technical skills: {profile.technical_skills}")
    # print(f"Company experience: {profile.company_experience}")
    
    # Test mit einem Verzeichnis
    # profiles = parser.parse_directory("path/to/profiles")
    # print(f"Parsed {len(profiles)} profiles")
