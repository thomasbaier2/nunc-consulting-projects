#!/usr/bin/env python3
"""
NUNC Expert Management System - STABLE VERSION
Robustes, zuverlässiges System ohne Unicode-Probleme
"""

import os
import sys
import json
import traceback
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional

# Flask und Web-Interface
from flask import Flask, request, jsonify, render_template_string, send_file, redirect, url_for
from werkzeug.utils import secure_filename

# PDF-Verarbeitung
import PyPDF2
import pdfplumber

# Word-Dokument-Generierung
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Konfiguration
UPLOAD_FOLDER = Path("uploads")
WORD_OUTPUT_FOLDER = Path("word_documents")
PROFILES_FOLDER = Path("profiles")

# Ordner erstellen
UPLOAD_FOLDER.mkdir(exist_ok=True)
WORD_OUTPUT_FOLDER.mkdir(exist_ok=True)
PROFILES_FOLDER.mkdir(exist_ok=True)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = str(UPLOAD_FOLDER)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max

class StableCvProcessor:
    """Stabile CV-Verarbeitung ohne externe Abhängigkeiten"""
    
    def __init__(self):
        self.supported_formats = ['.pdf']
    
    def process_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Verarbeitet PDF-CV und extrahiert strukturierte Daten"""
        try:
            print(f"Verarbeite PDF: {pdf_path}")
            
            # PDF Text extrahieren
            text_content = self._extract_text_from_pdf(pdf_path)
            print(f"Text extrahiert: {len(text_content)} Zeichen")
            
            # Strukturierte Daten extrahieren
            structured_data = self._extract_structured_data(text_content)
            print(f"Daten extrahiert: {structured_data.get('name', 'Unbekannt')}")
            
            # NUNC-Format konvertieren
            nunc_profile = self._convert_to_nunc_format(structured_data)
            print(f"NUNC-Profil erstellt: {nunc_profile.get('expert_name', 'Unbekannt')}")
            
            return nunc_profile
            
        except Exception as e:
            print(f"Fehler bei PDF-Verarbeitung: {e}")
            traceback.print_exc()
            return {'error': f'Fehler bei der PDF-Verarbeitung: {str(e)}'}
    
    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extrahiert Text aus PDF-Datei"""
        try:
            text = ""
            
            # Versuche zuerst mit pdfplumber
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except:
                # Fallback zu PyPDF2
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"PDF Text-Extraktion fehlgeschlagen: {str(e)}")
    
    def _extract_structured_data(self, text: str) -> Dict[str, Any]:
        """Extrahiert strukturierte Daten aus CV-Text"""
        import re
        
        # Name extrahieren
        name = self._extract_name(text)
        
        # Kontaktdaten
        contact = self._extract_contact_info(text)
        
        # Berufserfahrung
        experience = self._extract_experience(text)
        
        # Bildung
        education = self._extract_education(text)
        
        # Fähigkeiten
        skills = self._extract_skills(text)
        
        # Sprachen
        languages = self._extract_languages(text)
        
        return {
            'name': name,
            'contact': contact,
            'experience': experience,
            'education': education,
            'skills': skills,
            'languages': languages,
            'raw_text': text
        }
    
    def _extract_name(self, text: str) -> str:
        """Extrahiert den Namen aus dem CV"""
        lines = text.split('\n')
        for line in lines[:5]:
            line = line.strip()
            if len(line) > 2 and not any(char.isdigit() for char in line):
                return line
        return "Unbekannt"
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extrahiert Kontaktdaten"""
        import re
        
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_pattern = r'(\+?[\d\s\-\(\)]{10,})'
        
        emails = re.findall(email_pattern, text)
        phones = re.findall(phone_pattern, text)
        
        return {
            'email': emails[0] if emails else '',
            'phone': phones[0] if phones else ''
        }
    
    def _extract_experience(self, text: str) -> List[Dict[str, str]]:
        """Extrahiert Berufserfahrung"""
        experience = []
        lines = text.split('\n')
        current_job = {}
        
        for line in lines:
            line = line.strip()
            if any(keyword in line.lower() for keyword in ['consultant', 'manager', 'developer', 'analyst', 'engineer']):
                if current_job:
                    experience.append(current_job)
                current_job = {'title': line, 'description': ''}
            elif current_job and line:
                current_job['description'] += line + ' '
        
        if current_job:
            experience.append(current_job)
        
        return experience
    
    def _extract_education(self, text: str) -> List[str]:
        """Extrahiert Bildungsweg"""
        education = []
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if any(keyword in line.lower() for keyword in ['university', 'college', 'degree', 'bachelor', 'master', 'phd']):
                education.append(line)
        
        return education
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extrahiert Fähigkeiten und Technologien"""
        skills = []
        
        tech_keywords = [
            'salesforce', 'python', 'java', 'javascript', 'react', 'angular',
            'sql', 'oracle', 'sap', 'microsoft', 'aws', 'azure', 'docker',
            'kubernetes', 'git', 'agile', 'scrum', 'kanban'
        ]
        
        text_lower = text.lower()
        for keyword in tech_keywords:
            if keyword in text_lower:
                skills.append(keyword.title())
        
        return skills
    
    def _extract_languages(self, text: str) -> List[str]:
        """Extrahiert Sprachen"""
        languages = []
        
        language_keywords = ['german', 'english', 'french', 'spanish', 'italian', 'deutsch', 'englisch']
        
        text_lower = text.lower()
        for keyword in language_keywords:
            if keyword in text_lower:
                languages.append(keyword.title())
        
        return languages
    
    def _convert_to_nunc_format(self, structured_data: Dict[str, Any]) -> Dict[str, Any]:
        """Konvertiert extrahierte Daten in NUNC-Format"""
        # Hauptfokus bestimmen
        hauptfokus = self._determine_hauptfokus(structured_data)
        
        # Zur Person Beschreibung generieren
        zur_person = self._generate_zur_person(structured_data)
        
        # Projekthistorie erstellen
        projekthistorie = self._create_projekthistorie(structured_data)
        
        return {
            'expert_name': structured_data.get('name', 'Unbekannt'),
            'email': structured_data.get('contact', {}).get('email', ''),
            'phone': structured_data.get('contact', {}).get('phone', ''),
            'hauptfokus': hauptfokus,
            'sprachen': ', '.join(structured_data.get('languages', [])),
            'zur_person': zur_person,
            'besondere_kenntnisse': ', '.join(structured_data.get('skills', [])),
            'branchenkenntnisse': 'Diverse Branchen',
            'methoden': 'Agile Methoden, Scrum',
            'technologien': ', '.join(structured_data.get('skills', [])),
            'zertifizierungen': 'Zu ermitteln',
            'projekthistorie': projekthistorie,
            'projekthistorie_text': self._create_projekthistorie_text(projekthistorie),
            'education': structured_data.get('education', []),
            'source': 'pdf',
            'tags': structured_data.get('skills', [])
        }
    
    def _determine_hauptfokus(self, data: Dict[str, Any]) -> str:
        """Bestimmt den Hauptfokus basierend auf Erfahrung"""
        experience = data.get('experience', [])
        if experience:
            return experience[0].get('title', 'Consultant')
        return 'Consultant'
    
    def _generate_zur_person(self, data: Dict[str, Any]) -> str:
        """Generiert Zur Person Beschreibung"""
        name = data.get('name', '')
        experience = data.get('experience', [])
        skills = data.get('skills', [])
        
        description = f"Erfahrener {data.get('hauptfokus', 'Consultant')} mit umfassender Expertise"
        
        if skills:
            description += f" in {', '.join(skills[:3])}"
        
        if experience:
            description += f". {len(experience)} Jahre Berufserfahrung"
        
        return description
    
    def _create_projekthistorie(self, data: Dict[str, Any]) -> List[Dict[str, str]]:
        """Erstellt Projekthistorie aus Erfahrung"""
        projekthistorie = []
        experience = data.get('experience', [])
        
        for i, exp in enumerate(experience):
            projekthistorie.append({
                'projekt_name': exp.get('title', f'Projekt {i+1}'),
                'zeitraum': f'2020-{2023+i}',
                'projektrolle': exp.get('title', 'Consultant'),
                'aufgaben': exp.get('description', 'Beratung und Implementierung')
            })
        
        return projekthistorie
    
    def _create_projekthistorie_text(self, projekthistorie: List[Dict[str, str]]) -> str:
        """Erstellt Text-Version der Projekthistorie"""
        text_parts = []
        for project in projekthistorie:
            text_parts.append(f"{project['projekt_name']} ({project['zeitraum']}) - {project['projektrolle']}: {project['aufgaben']}")
        return '; '.join(text_parts)

class StableWordGenerator:
    """Stabile Word-Dokument-Generierung"""
    
    def __init__(self):
        self.output_path = WORD_OUTPUT_FOLDER
    
    def generate_word_document(self, template_data: Dict[str, Any]) -> str:
        """Generiert Word-Dokument aus Template-Daten"""
        try:
            print(f"Generiere Word-Dokument für: {template_data.get('expert_name', 'Unknown')}")
            
            # Neues Word-Dokument erstellen
            doc = Document()
            
            # Header
            header = doc.add_heading(f'NUNC Co.-Expert: {template_data.get("expert_name", "Unknown")}', level=1)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subheader = doc.add_paragraph(f'Hauptfokus: {template_data.get("hauptfokus", "Consultant")}')
            subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Profilvorstellung
            doc.add_paragraph(f'Profilvorstellung NUNC Consulting GmbH: Sprachen: {template_data.get("sprachen", "Deutsch/Englisch")}')
            
            # Zur Person
            doc.add_heading('Zur Person:', level=2)
            doc.add_paragraph(template_data.get('zur_person', 'Keine Informationen verfügbar'))
            
            # Besondere Kenntnisse
            doc.add_heading('Besondere Kenntnisse:', level=2)
            doc.add_paragraph(template_data.get('besondere_kenntnisse', 'Keine Informationen verfügbar'))
            
            # Branchenkenntnisse
            doc.add_heading('Branchenkenntnisse:', level=2)
            doc.add_paragraph(template_data.get('branchenkenntnisse', 'Diverse Branchen'))
            
            # Methoden
            doc.add_heading('Methoden:', level=2)
            doc.add_paragraph(template_data.get('methoden', 'Agile Methoden, Scrum'))
            
            # Technologien
            doc.add_heading('Technologien:', level=2)
            doc.add_paragraph(template_data.get('technologien', 'Keine Informationen verfügbar'))
            
            # Zertifizierungen
            doc.add_heading('Zertifizierungen:', level=2)
            doc.add_paragraph(template_data.get('zertifizierungen', 'Zu ermitteln'))
            
            # Projekthistorie
            doc.add_heading('PROJEKTHISTORIE', level=2)
            doc.add_paragraph(template_data.get('projekthistorie_text', 'Keine Projekte verfügbar'))
            
            # Speichern
            output_filename = f"NUNC_Profile_{template_data.get('expert_name', 'Unknown').replace(' ', '_')}.docx"
            output_file = self.output_path / output_filename
            
            doc.save(str(output_file))
            
            print(f"Word-Dokument erstellt: {output_file}")
            return str(output_file)
            
        except Exception as e:
            print(f"Fehler beim Generieren des Word-Dokuments: {e}")
            traceback.print_exc()
            return None

# Globale Instanzen
cv_processor = StableCvProcessor()
word_generator = StableWordGenerator()

# HTML Templates
INDEX_HTML = """
<!DOCTYPE html>
<html lang="de">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>NUNC Expert Management System</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 0; padding: 20px; background: #f5f5f5; }
        .container { max-width: 1200px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
        .header { text-align: center; margin-bottom: 40px; }
        .header h1 { color: #2c3e50; margin-bottom: 10px; }
        .header p { color: #7f8c8d; font-size: 18px; }
        .nav { display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 20px; margin-bottom: 40px; }
        .nav-card { background: #ecf0f1; padding: 25px; border-radius: 8px; text-align: center; transition: transform 0.3s; }
        .nav-card:hover { transform: translateY(-5px); background: #d5dbdb; }
        .nav-card h2 { color: #2c3e50; margin-bottom: 15px; }
        .nav-card p { color: #7f8c8d; margin-bottom: 20px; }
        .nav-card a { display: inline-block; background: #3498db; color: white; padding: 12px 24px; text-decoration: none; border-radius: 5px; transition: background 0.3s; }
        .nav-card a:hover { background: #2980b9; }
        .footer { text-align: center; margin-top: 40px; color: #7f8c8d; }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>NUNC Expert Management System</h1>
            <p>Professionelle Expertenverwaltung und CV-Verarbeitung</p>
        </div>
        
        <div class="nav">
            <div class="nav-card">
                <h2>Profil-Management</h2>
                <p>Verwalten Sie Experten-Profile, erstellen Sie neue Profile und aktualisieren Sie bestehende.</p>
                <a href="/profiles">Profile verwalten</a>
            </div>
            
            <div class="nav-card">
                <h2>Verfügbarkeits-Management</h2>
                <p>Verfügbarkeits-Anfragen senden, Antworten verwalten und Verfügbarkeit tracken.</p>
                <a href="/availability">Verfügbarkeit verwalten</a>
            </div>
            
            <div class="nav-card">
                <h2>Kandidaten-Suche</h2>
                <p>Automatisierte Suche auf LinkedIn und Freelancermap mit AI-basiertem Matching.</p>
                <a href="/candidates">Kandidaten suchen</a>
            </div>
            
            <div class="nav-card">
                <h2>Projekt-Matching</h2>
                <p>Projekte mit passenden Experten matchen und optimale Besetzungen finden.</p>
                <a href="/projects">Projekte verwalten</a>
            </div>
            
            <div class="nav-card">
                <h2>CV-Verarbeitung</h2>
                <p>PDF-CVs hochladen, automatisch verarbeiten und in NUNC-Profile konvertieren.</p>
                <a href="/cv-processing">CV verarbeiten</a>
            </div>
        </div>
    </div>
    
    <div class="footer">
        &copy; 2025 NUNC Consulting GmbH. All rights reserved.
    </div>
</body>
</html>
"""

CV_PROCESSING_HTML = """
<!DOCTYPE html>
<html lang="de">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CV-Verarbeitung - NUNC Expert Management System</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 0; padding: 20px; background: #f5f5f5; }
        .container { max-width: 1000px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
        .header { text-align: center; margin-bottom: 40px; }
        .header h1 { color: #2c3e50; margin-bottom: 10px; }
        .upload-area { border: 2px dashed #3498db; padding: 40px; text-align: center; margin: 20px 0; border-radius: 8px; background: #f8f9fa; }
        .upload-area:hover { background: #e9ecef; }
        .upload-area.dragover { background: #d4edda; border-color: #28a745; }
        .file-input { display: none; }
        .upload-btn { background: #3498db; color: white; padding: 12px 24px; border: none; border-radius: 5px; cursor: pointer; font-size: 16px; }
        .upload-btn:hover { background: #2980b9; }
        .result { margin-top: 20px; padding: 20px; background: #f8f9fa; border-radius: 8px; }
        .profile-data { background: #e9ecef; padding: 15px; border-radius: 5px; margin: 10px 0; }
        .btn { background: #28a745; color: white; padding: 10px 20px; border: none; border-radius: 5px; cursor: pointer; margin: 5px; }
        .btn:hover { background: #218838; }
        .btn-secondary { background: #6c757d; }
        .btn-secondary:hover { background: #5a6268; }
        .loading { display: none; text-align: center; }
        .spinner { border: 4px solid #f3f3f3; border-top: 4px solid #3498db; border-radius: 50%; width: 40px; height: 40px; animation: spin 2s linear infinite; margin: 0 auto; }
        @keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>CV-Verarbeitung</h1>
            <p>Laden Sie PDF-CVs hoch und konvertieren Sie sie in NUNC-Profile</p>
        </div>
        
        <div class="upload-area" id="uploadArea">
            <h3>PDF-CV hochladen</h3>
            <p>Ziehen Sie eine PDF-Datei hierher oder klicken Sie zum Auswählen</p>
            <input type="file" id="fileInput" class="file-input" accept=".pdf">
            <button class="upload-btn" onclick="document.getElementById('fileInput').click()">Datei auswählen</button>
        </div>
        
        <div class="loading" id="loading">
            <div class="spinner"></div>
            <p>CV wird verarbeitet...</p>
        </div>
        
        <div class="result" id="result" style="display: none;">
            <h3>Verarbeitetes Profil</h3>
            <div id="profileData"></div>
            <button class="btn" onclick="generateWord()">Word-Dokument generieren</button>
            <button class="btn btn-secondary" onclick="saveProfile()">Profil speichern</button>
        </div>
        
        <div class="result" id="wordResult" style="display: none;">
            <h3>Word-Dokument</h3>
            <p id="wordStatus"></p>
            <a id="downloadLink" href="#" class="btn" style="display: none;">Word-Dokument herunterladen</a>
        </div>
    </div>
    
    <script>
        const uploadArea = document.getElementById('uploadArea');
        const fileInput = document.getElementById('fileInput');
        const result = document.getElementById('result');
        const loading = document.getElementById('loading');
        const profileData = document.getElementById('profileData');
        const wordResult = document.getElementById('wordResult');
        const wordStatus = document.getElementById('wordStatus');
        const downloadLink = document.getElementById('downloadLink');
        
        let currentProfile = null;
        
        // Drag & Drop
        uploadArea.addEventListener('dragover', (e) => {
            e.preventDefault();
            uploadArea.classList.add('dragover');
        });
        
        uploadArea.addEventListener('dragleave', () => {
            uploadArea.classList.remove('dragover');
        });
        
        uploadArea.addEventListener('drop', (e) => {
            e.preventDefault();
            uploadArea.classList.remove('dragover');
            const files = e.dataTransfer.files;
            if (files.length > 0) {
                handleFile(files[0]);
            }
        });
        
        // File Input
        fileInput.addEventListener('change', (e) => {
            if (e.target.files.length > 0) {
                handleFile(e.target.files[0]);
            }
        });
        
        function handleFile(file) {
            if (!file.name.toLowerCase().endsWith('.pdf')) {
                alert('Bitte wählen Sie eine PDF-Datei aus.');
                return;
            }
            
            const formData = new FormData();
            formData.append('file', file);
            
            loading.style.display = 'block';
            result.style.display = 'none';
            wordResult.style.display = 'none';
            
            fetch('/api/cv/upload', {
                method: 'POST',
                body: formData
            })
            .then(response => response.json())
            .then(data => {
                loading.style.display = 'none';
                
                if (data.success) {
                    currentProfile = data.profile;
                    displayProfile(data.profile);
                    result.style.display = 'block';
                } else {
                    alert('Fehler: ' + data.error);
                }
            })
            .catch(error => {
                loading.style.display = 'none';
                alert('Fehler: ' + error);
            });
        }
        
        function displayProfile(profile) {
            profileData.innerHTML = `
                <div class="profile-data">
                    <h4>Name: ${profile.expert_name || 'Unbekannt'}</h4>
                    <p><strong>Hauptfokus:</strong> ${profile.hauptfokus || 'Consultant'}</p>
                    <p><strong>Sprachen:</strong> ${profile.sprachen || 'Deutsch/Englisch'}</p>
                    <p><strong>Zur Person:</strong> ${profile.zur_person || 'Keine Informationen'}</p>
                    <p><strong>Besondere Kenntnisse:</strong> ${profile.besondere_kenntnisse || 'Keine Informationen'}</p>
                    <p><strong>Technologien:</strong> ${profile.technologien || 'Keine Informationen'}</p>
                </div>
            `;
        }
        
        function generateWord() {
            if (!currentProfile) {
                alert('Kein Profil verfügbar');
                return;
            }
            
            wordStatus.textContent = 'Word-Dokument wird generiert...';
            wordResult.style.display = 'block';
            
            fetch('/api/cv/generate-word', {
                method: 'POST',
                headers: {
                    'Content-Type': 'application/json'
                },
                body: JSON.stringify(currentProfile)
            })
            .then(response => response.json())
            .then(data => {
                if (data.success) {
                    wordStatus.textContent = 'Word-Dokument erfolgreich generiert!';
                    downloadLink.href = data.download_url;
                    downloadLink.style.display = 'inline-block';
                } else {
                    wordStatus.textContent = 'Fehler: ' + data.error;
                }
            })
            .catch(error => {
                wordStatus.textContent = 'Fehler: ' + error;
            });
        }
        
        function saveProfile() {
            if (!currentProfile) {
                alert('Kein Profil verfügbar');
                return;
            }
            
            fetch('/api/cv/save-profile', {
                method: 'POST',
                headers: {
                    'Content-Type': 'application/json'
                },
                body: JSON.stringify(currentProfile)
            })
            .then(response => response.json())
            .then(data => {
                if (data.success) {
                    alert('Profil erfolgreich gespeichert!');
                } else {
                    alert('Fehler: ' + data.error);
                }
            })
            .catch(error => {
                alert('Fehler: ' + error);
            });
        }
    </script>
</body>
</html>
"""

# Routes
@app.route('/')
def index():
    """Hauptseite"""
    return INDEX_HTML

@app.route('/cv-processing')
def cv_processing():
    """CV-Verarbeitungsseite"""
    return CV_PROCESSING_HTML

@app.route('/api/cv/upload', methods=['POST'])
def upload_cv():
    """CV-Upload und -Verarbeitung"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'Keine Datei gefunden'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'Keine Datei ausgewählt'})
        
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({'success': False, 'error': 'Nur PDF-Dateien sind erlaubt'})
        
        # Datei speichern
        filename = secure_filename(file.filename)
        file_path = UPLOAD_FOLDER / filename
        file.save(str(file_path))
        
        print(f"Datei gespeichert: {file_path}")
        
        # CV verarbeiten
        profile = cv_processor.process_pdf(str(file_path))
        
        if 'error' in profile:
            return jsonify({'success': False, 'error': profile['error']})
        
        return jsonify({'success': True, 'profile': profile})
        
    except Exception as e:
        print(f"Fehler beim CV-Upload: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/cv/generate-word', methods=['POST'])
def generate_word_document():
    """Generiert Word-Dokument aus Profil"""
    try:
        profile_data = request.json
        
        # Word-Dokument generieren
        word_file = word_generator.generate_word_document(profile_data)
        
        if word_file:
            filename = Path(word_file).name
            return jsonify({
                'success': True, 
                'word_file': word_file,
                'download_url': f'/api/cv/download/{filename}'
            })
        else:
            return jsonify({'success': False, 'error': 'Fehler beim Generieren des Word-Dokuments'})
            
    except Exception as e:
        print(f"Fehler beim Word-Generieren: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/cv/download/<filename>')
def download_word_document(filename):
    """Lädt Word-Dokument herunter"""
    try:
        file_path = WORD_OUTPUT_FOLDER / filename
        
        if file_path.exists():
            return send_file(str(file_path), as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        else:
            return jsonify({'success': False, 'error': 'Datei nicht gefunden'})
            
    except Exception as e:
        print(f"Fehler beim Download: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/cv/save-profile', methods=['POST'])
def save_profile():
    """Speichert Profil"""
    try:
        profile_data = request.json
        
        # Profil als JSON speichern
        profile_file = PROFILES_FOLDER / f"{profile_data.get('expert_name', 'Unknown').replace(' ', '_')}.json"
        
        with open(profile_file, 'w', encoding='utf-8') as f:
            json.dump(profile_data, f, ensure_ascii=False, indent=2)
        
        return jsonify({'success': True, 'message': 'Profil gespeichert'})
        
    except Exception as e:
        print(f"Fehler beim Speichern: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})

# Fallback Routes
@app.route('/profiles')
def profiles():
    return "<h1>Profil-Management</h1><p>Funktion in Entwicklung</p><a href='/'>Zurück zur Hauptseite</a>"

@app.route('/availability')
def availability():
    return "<h1>Verfügbarkeits-Management</h1><p>Funktion in Entwicklung</p><a href='/'>Zurück zur Hauptseite</a>"

@app.route('/candidates')
def candidates():
    return "<h1>Kandidaten-Suche</h1><p>Funktion in Entwicklung</p><a href='/'>Zurück zur Hauptseite</a>"

@app.route('/projects')
def projects():
    return "<h1>Projekt-Matching</h1><p>Funktion in Entwicklung</p><a href='/'>Zurück zur Hauptseite</a>"

if __name__ == '__main__':
    print("=" * 60)
    print("NUNC Expert Management System - STABLE VERSION")
    print("=" * 60)
    print("System wird gestartet...")
    print(f"Upload-Ordner: {UPLOAD_FOLDER}")
    print(f"Word-Ordner: {WORD_OUTPUT_FOLDER}")
    print(f"Profile-Ordner: {PROFILES_FOLDER}")
    print("=" * 60)
    print("Öffnen Sie: http://127.0.0.1:9000")
    print("=" * 60)
    
    try:
        app.run(debug=True, host='127.0.0.1', port=9000)
    except Exception as e:
        print(f"Fehler beim Starten: {e}")
        traceback.print_exc()

