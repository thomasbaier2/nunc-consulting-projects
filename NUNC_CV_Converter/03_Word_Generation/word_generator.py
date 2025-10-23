#!/usr/bin/env python3
"""
NUNC CV Converter - Word Document Generator
Erstellt echte .docx Dateien mit docxtemplater
"""

import os
import json
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any
import base64

try:
    from docxtpl import DocxTemplate
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("Installing required packages...")
    os.system("pip3 install python-docx docxtpl")
    from docxtpl import DocxTemplate
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

class NuncWordGenerator:
    def __init__(self):
        self.template_path = "../07_Output_Files/word_documents/nunc_template.docx"
        self.output_path = "../07_Output_Files/word_documents"
        
    def create_nunc_template(self) -> str:
        """Erstellt ein NUNC Word-Template mit Platzhaltern"""
        doc = Document()
        
        # Header
        header = doc.add_heading('NUNC Co.-Expert: {{expert_name}}', level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subheader = doc.add_paragraph('Hauptfokus: {{hauptfokus}}')
        subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Profilvorstellung
        doc.add_paragraph('Profilvorstellung NUNC Consulting GmbH: Sprachen: {{sprachen}}')
        
        # Zur Person
        doc.add_heading('Zur Person:', level=2)
        doc.add_paragraph('{{zur_person}}')
        
        # Besondere Kenntnisse
        doc.add_heading('Besondere Kenntnisse:', level=2)
        doc.add_paragraph('{{besondere_kenntnisse}}')
        
        # Branchenkenntnisse
        doc.add_heading('Branchenkenntnisse:', level=2)
        doc.add_paragraph('{{branchenkenntnisse}}')
        
        # Methoden
        doc.add_heading('Methoden:', level=2)
        doc.add_paragraph('{{methoden}}')
        
        # Technologien
        doc.add_heading('Technologien:', level=2)
        doc.add_paragraph('{{technologien}}')
        
        # Zertifizierungen
        doc.add_heading('Zertifizierungen:', level=2)
        doc.add_paragraph('{{zertifizierungen}}')
        
        # Projekthistorie
        doc.add_heading('PROJEKTHISTORIE', level=2)
        
        # Projekte als strukturierte Liste
        doc.add_paragraph('Projekthistorie:')
        doc.add_paragraph('{{projekthistorie_text}}')
        
        # Template speichern
        os.makedirs(os.path.dirname(self.template_path), exist_ok=True)
        doc.save(self.template_path)
        
        return self.template_path
    
    def generate_word_document(self, template_data: Dict[str, Any]) -> str:
        """Generiert Word-Dokument aus Template-Daten"""
        try:
            # Template erstellen falls nicht vorhanden
            if not os.path.exists(self.template_path):
                self.create_nunc_template()
            
            # Template laden
            doc = DocxTemplate(self.template_path)
            
            # Daten in Template einsetzen
            doc.render(template_data)
            
            # Output-Datei generieren
            output_filename = f"NUNC_Profile_{template_data['expert_name'].replace(' ', '_')}.docx"
            output_file = os.path.join(self.output_path, output_filename)
            
            os.makedirs(self.output_path, exist_ok=True)
            doc.save(output_file)
            
            return output_file
            
        except Exception as e:
            print(f"Fehler beim Generieren des Word-Dokuments: {e}")
            return None

def main():
    """Test der Word-Generierung"""
    print("🚀 NUNC Word Generator Test")
    
    generator = NuncWordGenerator()
    
    # Template erstellen
    template_path = generator.create_nunc_template()
    print(f"✅ Template erstellt: {template_path}")
    
    # Test-Daten
    test_data = {
        'expert_name': 'Lukas Pfanner',
        'hauptfokus': 'Salesforce Consultant',
        'sprachen': 'Deutsch/Englisch',
        'zur_person': 'Erfahrener Salesforce Consultant mit umfassender Expertise in verschiedenen Technologien.',
        'besondere_kenntnisse': 'Zertifizierungen in Salesforce und Projektmanagement.',
        'branchenkenntnisse': 'Telecommunications/Retail/Oil & Energy',
        'methoden': 'Agile Methoden, Scrum, PRINCE2',
        'technologien': 'Salesforce, CRM, Projektmanagement',
        'zertifizierungen': 'Salesforce Certified Administrator',
        'projekthistorie_text': 'Test Projekt (2023-2024) - Consultant: Beratung und Implementierung'
    }
    
    # Word-Dokument generieren
    output_file = generator.generate_word_document(test_data)
    
    if output_file:
        print(f"✅ Word-Dokument erstellt: {output_file}")
    else:
        print("❌ Fehler beim Erstellen des Word-Dokuments")

if __name__ == "__main__":
    main()

